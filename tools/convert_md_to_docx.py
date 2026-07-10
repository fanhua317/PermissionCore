"""Convert a Markdown file to a simple editable DOCX document.

Dependency: ``python -m pip install python-docx``
Usage: ``python tools/convert_md_to_docx.py INPUT.md [-o OUTPUT.docx] [--force]``
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError as exc:  # pragma: no cover - depends on the local runtime
    raise SystemExit(
        "Missing dependency 'python-docx'. Install it with: "
        "python -m pip install python-docx"
    ) from exc


BODY_FONT = "SimSun"
BODY_SIZE_PT = 12
LINE_SPACING = 1.25
HEADING_SIZES = {1: 20, 2: 16, 3: 14, 4: 12, 5: 12, 6: 12}
TABLE_SEPARATOR = re.compile(
    r"^\s*\|?\s*:?-{3,}:?\s*(?:\|\s*:?-{3,}:?\s*)+\|?\s*$"
)


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input", type=Path, help="source Markdown file")
    parser.add_argument(
        "-o", "--output", type=Path, help="output DOCX path (default: beside input)"
    )
    parser.add_argument("-f", "--force", action="store_true", help="overwrite output")
    return parser.parse_args()


def set_run_font(
    run, name: str = BODY_FONT, size_pt: int = BODY_SIZE_PT,
    bold: bool = False, italic: bool = False,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic


def set_paragraph_format(paragraph) -> None:
    paragraph.paragraph_format.line_spacing = LINE_SPACING


def add_heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    set_run_font(
        paragraph.add_run(text), size_pt=HEADING_SIZES.get(level, 12), bold=True
    )
    set_paragraph_format(paragraph)


def add_paragraph(doc, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_run_font(paragraph.add_run(text))
    set_paragraph_format(paragraph)


def add_code_block(doc, lines: list[str]) -> None:
    for line in lines or [""]:
        paragraph = doc.add_paragraph()
        set_run_font(paragraph.add_run(line), name="Courier New", size_pt=10)
        set_paragraph_format(paragraph)


def split_table_row(line: str) -> list[str]:
    cells = [cell.strip() for cell in line.strip().split("|")]
    if cells and not cells[0]:
        cells.pop(0)
    if cells and not cells[-1]:
        cells.pop()
    return cells


def add_table(doc, block: list[str]) -> None:
    rows = [split_table_row(line) for line in block]
    if len(rows) > 1 and TABLE_SEPARATOR.match(block[1]):
        rows.pop(1)
    column_count = max((len(row) for row in rows), default=0)
    if not column_count:
        return
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.rows[row_index].cells[column_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            set_run_font(paragraph.add_run(value))
            set_paragraph_format(paragraph)


def convert(source: Path, output: Path) -> None:
    if not source.is_file():
        raise FileNotFoundError(f"Markdown file not found: {source}")
    if output.suffix.lower() != ".docx":
        raise ValueError("Output path must end with .docx")

    lines = source.read_text(encoding="utf-8-sig").splitlines()
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = section.right_margin = Cm(2.54)
    section.top_margin = section.bottom_margin = Cm(2.54)

    index = 0
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if (
            "|" in line
            and index + 1 < len(lines)
            and TABLE_SEPARATOR.match(lines[index + 1])
        ):
            block = [line, lines[index + 1]]
            index += 2
            while index < len(lines) and "|" in lines[index] and lines[index].strip():
                block.append(lines[index])
                index += 1
            add_table(doc, block)
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            add_heading(doc, heading.group(2).strip(), len(heading.group(1)))
        elif match := re.match(r"^[*+-]\s+(.*)$", line):
            paragraph = doc.add_paragraph(style="List Bullet")
            set_run_font(paragraph.add_run(match.group(1).strip()))
            set_paragraph_format(paragraph)
        elif match := re.match(r"^\d+\.\s+(.*)$", line):
            paragraph = doc.add_paragraph(style="List Number")
            set_run_font(paragraph.add_run(match.group(1).strip()))
            set_paragraph_format(paragraph)
        elif line.strip():
            add_paragraph(doc, line)
        else:
            doc.add_paragraph("")
        index += 1

    if in_code:
        add_code_block(doc, code_lines)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> None:
    args = parse_args()
    source = args.input.expanduser().resolve()
    output = (args.output or source.with_suffix(".docx")).expanduser().resolve()
    if source == output:
        raise ValueError("Input and output paths must be different")
    if output.exists() and not args.force:
        raise SystemExit(f"Output already exists (use --force to overwrite): {output}")
    try:
        convert(source, output)
    except (FileNotFoundError, UnicodeError, ValueError) as exc:
        raise SystemExit(str(exc)) from exc
    print(f"Saved DOCX to: {output}")


if __name__ == "__main__":
    main()
